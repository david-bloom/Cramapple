import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "architecture" / "HIGH_LEVEL_SYSTEM_ARCHITECTURE.md"
OUTPUT = ROOT / "docs" / "architecture" / "Cramapple-High-Level-System-Architecture-v0.1.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 58)
MUTED = RGBColor(90, 99, 108)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def create_numbering(doc, ordered, start_at=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_at))
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_inline_runs(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    apply_numbering(paragraph, num_id)
    add_inline_runs(paragraph, text)


def add_code_block(doc, lines, keep_together=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    if keep_together:
        prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=100, start=180, bottom=100, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK, italic=True)


def table_widths(column_count):
    if column_count == 2:
        return [3000, 6360]
    if column_count == 3:
        return [2600, 4200, 2560]
    if column_count == 4:
        return [1900, 3000, 2300, 2160]
    if column_count == 5:
        return [2200, 1200, 1200, 1700, 3060]
    base = TABLE_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc, rows):
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(column_count))
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(rows):
        prevent_row_split(table.rows[row_index])
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline_runs(paragraph, text)
            for run in paragraph.runs:
                set_run_font(run, size=9.2, color=INK, bold=(row_index == 0))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def configure_document(header_text):
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = header_text
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    add_page_number(section.footer.paragraphs[0])
    return doc


def add_masthead(doc, title_text, version, date_text, status, scope):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(title_text)
    set_run_font(run, size=25, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(f"Canonical planning draft | Version {version}")
    set_run_font(run, size=13, color=MUTED)

    metadata = [
        ("Product owner", "David Bloom"),
        ("Date", date_text),
        ("Status", status),
        ("Scope", scope),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def render_markdown(doc, lines):
    bullet_id = create_numbering(doc, ordered=False)
    numbered_id = None
    in_code = False
    code_language = ""
    code_lines = []
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("```"):
            numbered_id = None
            if in_code:
                add_code_block(
                    doc,
                    code_lines,
                    keep_together=(code_language == "mermaid" and len(code_lines) <= 45),
                )
                code_lines = []
                code_language = ""
                in_code = False
            else:
                in_code = True
                code_language = line[3:].strip()
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not line:
            numbered_id = None
            index += 1
            continue

        if line.startswith("# "):
            numbered_id = None
            index += 1
            continue

        if line.startswith("**Canonical planning draft"):
            numbered_id = None
            index += 1
            continue

        if line.startswith("## "):
            numbered_id = None
            doc.add_paragraph(line[3:], style="Heading 1")
            index += 1
            continue

        if line.startswith("### "):
            numbered_id = None
            doc.add_paragraph(line[4:], style="Heading 2")
            index += 1
            continue

        if line.startswith("#### "):
            numbered_id = None
            doc.add_paragraph(line[5:], style="Heading 3")
            index += 1
            continue

        if line.startswith("|"):
            numbered_id = None
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        if line.startswith("> "):
            numbered_id = None
            add_quote(doc, line[2:])
            index += 1
            continue

        bullet_match = re.match(r"^- (.+)$", line)
        if bullet_match:
            numbered_id = None
            add_list_item(doc, bullet_match.group(1), bullet_id)
            index += 1
            continue

        number_match = re.match(r"^(\d+)\. (.+)$", line)
        if number_match:
            if numbered_id is None:
                numbered_id = create_numbering(
                    doc,
                    ordered=True,
                    start_at=int(number_match.group(1)),
                )
            add_list_item(doc, number_match.group(2), numbered_id)
            index += 1
            continue

        numbered_id = None
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline_runs(paragraph, line)
        index += 1


def build(source, output, title, header, version, date_text, status, scope, keywords):
    doc = configure_document(header)
    add_masthead(doc, title, version, date_text, status, scope)
    render_markdown(doc, source.read_text(encoding="utf-8").splitlines())

    properties = doc.core_properties
    properties.title = title
    properties.subject = "Canonical planning draft"
    properties.author = "Cramapple"
    properties.keywords = keywords

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build a Cramapple planning DOCX from Markdown.")
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    parser.add_argument("--title", default="Cramapple High-Level System Architecture")
    parser.add_argument("--header", default="CRAMAPPLE | HIGH-LEVEL SYSTEM ARCHITECTURE")
    parser.add_argument("--version", default="0.1")
    parser.add_argument("--date", dest="date_text", default="June 9, 2026")
    parser.add_argument("--status", default="Ready for owner review")
    parser.add_argument("--scope", default="Planning and architecture only; no implementation")
    parser.add_argument(
        "--keywords",
        default="Cramapple, architecture, AP exams, teaching, grading, validation",
    )
    args = parser.parse_args()
    build(
        source=args.source,
        output=args.output,
        title=args.title,
        header=args.header,
        version=args.version,
        date_text=args.date_text,
        status=args.status,
        scope=args.scope,
        keywords=args.keywords,
    )

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "product" / "CRAMAPPLE_VISION.md"
OUTPUT = ROOT / "docs" / "product" / "Cramapple-Vision-v0.3.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 45, 58)
MUTED = RGBColor(90, 99, 108)
LIGHT = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_hyperlink(paragraph, text, url, font_size=None):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    if font_size is not None:
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(int(font_size * 2)))
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(int(font_size * 2)))
        run_pr.extend([size, size_cs])
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    add_inline_runs(paragraph, text)


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    existing_num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    new_num_id = max(existing_num_ids, default=0) + 1

    style_num_id = doc.styles["List Number"].element.pPr.numPr.numId.val
    style_num = numbering.find(f".//{{{qn('w:num').split('}')[0][1:]}}}num[@{{{qn('w:numId').split('}')[0][1:]}}}numId='{style_num_id}']")
    abstract_num_id = style_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return new_num_id


def add_number(doc, text, num_id):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    add_inline_runs(paragraph, text)


def add_quote(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    table.rows[0].cells[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=140, start=220, bottom=140, end=220)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)

    subtitle = styles.add_style("Vision Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(16)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "CRAMAPPLE | VISION AND PROBLEM STATEMENT"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = MUTED

    add_page_number(section.footer.paragraphs[0])
    return doc


def render_markdown(doc, lines):
    in_refs = False
    first_heading = True
    current_num_id = None

    for raw in lines:
        line = raw.rstrip()
        if not line:
            current_num_id = None
            continue

        if line.startswith("# "):
            current_num_id = None
            if first_heading:
                p = doc.add_paragraph(style="Title")
                p.add_run(line[2:])
                first_heading = False
            continue

        if line.startswith("**Canonical reference"):
            current_num_id = None
            p = doc.add_paragraph(style="Vision Subtitle")
            add_inline_runs(p, line)
            continue

        if line.startswith("## "):
            current_num_id = None
            heading_text = line[3:]
            in_refs = heading_text == "Reference Sources"
            doc.add_paragraph(heading_text, style="Heading 1")
            continue

        if line.startswith("### "):
            current_num_id = None
            doc.add_paragraph(line[4:], style="Heading 2")
            continue

        if line.startswith("> "):
            current_num_id = None
            add_quote(doc, line[2:])
            continue

        bullet_match = re.match(r"^- (.+)$", line)
        if bullet_match:
            current_num_id = None
            text = bullet_match.group(1)
            if in_refs and ": http" in text:
                label, url = text.rsplit(": ", 1)
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                label_run = p.add_run(label + ": ")
                label_run.font.size = Pt(9.5)
                add_hyperlink(p, url, url, font_size=9.5)
            else:
                add_bullet(doc, text)
            continue

        number_match = re.match(r"^\d+\. (.+)$", line)
        if number_match:
            if current_num_id is None:
                current_num_id = create_numbering_instance(doc)
            add_number(doc, number_match.group(1), current_num_id)
            continue

        current_num_id = None
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline_runs(paragraph, line)


def build():
    doc = configure_document()
    render_markdown(doc, SOURCE.read_text(encoding="utf-8").splitlines())

    core_props = doc.core_properties
    core_props.title = "Cramapple Vision and Problem Statement"
    core_props.subject = "Canonical reference draft"
    core_props.author = "Cramapple"
    core_props.keywords = "Cramapple, AP Biology, score optimization, vision"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
